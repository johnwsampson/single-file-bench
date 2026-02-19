#!/usr/bin/env -S uv run --script
# /// script
# requires-python = ">=3.11"
# dependencies = [
#     "python-docx>=1.1.0",
#     "fastmcp",
# ]
# ///
"""Convert Microsoft Word (.docx) files to markdown.

Extract text content, tables, and structure from Word documents.
Outputs clean markdown suitable for piping to other tools.

Usage:
    sft_docx2md.py read document.docx
    sft_docx2md.py read --no-tables document.docx
    sft_docx2md.py analyze document.docx
    sft_docx2md.py read document.docx | sft_md2html.py --show
    sft_docx2md.py mcp-stdio
"""

import argparse
import json
import os
import sys
import time
from datetime import datetime, timezone
from pathlib import Path
from typing import Any


# =============================================================================
# LOGGING
# =============================================================================
_LEVELS = {"TRACE": 5, "DEBUG": 10, "INFO": 20, "WARN": 30, "ERROR": 40, "FATAL": 50}
_THRESHOLD = _LEVELS.get(os.environ.get("SFB_LOG_LEVEL", "INFO"), 20)
_LOG_DIR = os.environ.get("SFB_LOG_DIR", "")
_SCRIPT = Path(__file__).stem
_LOG = (
    Path(_LOG_DIR) / f"{_SCRIPT}_log.tsv"
    if _LOG_DIR
    else Path(__file__).parent / f"{_SCRIPT}_log.tsv"
)
_HEADER = "#timestamp\tscript\tlevel\tevent\tmessage\tdetail\tmetrics\ttrace\n"


def _log(
    level: str,
    event: str,
    msg: str,
    *,
    detail: str = "",
    metrics: str = "",
    trace: str = "",
):
    """Append TSV log line. Logging never crashes the main flow."""
    if _LEVELS.get(level, 20) < _THRESHOLD:
        return
    try:
        ts = datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%S")
        write_header = not _LOG.exists()
        with open(_LOG, "a") as f:
            if write_header:
                f.write(_HEADER)
            f.write(f"{ts}\t{_SCRIPT}\t{level}\t{event}\t{msg}\t{detail}\t{metrics}\t{trace}\n")
    except Exception:
        pass


# =============================================================================
# CONFIGURATION
# =============================================================================
EXPOSED = ["read", "analyze"]

CONFIG = {
    "version": "1.0.0",
}


# =============================================================================
# CORE FUNCTIONS
# =============================================================================

def _read_impl(filepath: str, include_tables: bool = True) -> tuple[str, dict]:
    """Read and convert Word document to markdown. CLI: read, MCP: read."""
    from docx import Document
    
    start_ms = time.time() * 1000
    
    path = Path(filepath)
    assert path.exists(), f"File not found: {filepath}"
    assert path.suffix.lower() == ".docx", f"Not a .docx file: {filepath}"
    
    document = Document(str(path))
    
    # Extract paragraphs
    content = []
    for para in document.paragraphs:
        text = para.text.strip()
        if text:
            # Check if it's a heading
            if para.style and para.style.name.startswith("Heading"):
                level_str = para.style.name.replace("Heading ", "").replace("Title", "0")
                try:
                    level = int(level_str)
                    prefix = "#" * (level + 1) + " "
                except ValueError:
                    prefix = "# "
                content.append(f"{prefix}{text}")
            else:
                content.append(text)
    
    # Extract tables if requested
    table_count = 0
    if include_tables and document.tables:
        for i, table in enumerate(document.tables):
            content.append(f"\n### Table {i + 1}")
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append("| " + " | ".join(cells) + " |")
            
            # Add header separator
            if rows:
                col_count = len(table.rows[0].cells)
                rows.insert(1, "| " + " | ".join(["---"] * col_count) + " |")
            
            content.extend(rows)
            content.append("")
            table_count += 1
    
    result = "\n\n".join(content)
    
    latency_ms = time.time() * 1000 - start_ms
    metrics = {
        "latency_ms": round(latency_ms, 2),
        "paragraphs": len(document.paragraphs),
        "tables": table_count,
    }
    
    return result, metrics


def _analyze_impl(filepath: str) -> tuple[dict, dict]:
    """Analyze Word document structure. CLI: analyze, MCP: analyze."""
    from docx import Document
    
    start_ms = time.time() * 1000
    
    path = Path(filepath)
    assert path.exists(), f"File not found: {filepath}"
    
    document = Document(str(path))
    
    # Style summary
    styles = {}
    for para in document.paragraphs:
        style_name = para.style.name if para.style else "Default"
        styles[style_name] = styles.get(style_name, 0) + 1
    
    # Table details
    table_info = []
    for i, table in enumerate(document.tables):
        rows = len(table.rows)
        cols = len(table.rows[0].cells) if rows > 0 else 0
        table_info.append({"index": i + 1, "rows": rows, "cols": cols})
    
    result = {
        "filename": path.name,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "styles": styles,
        "table_details": table_info,
    }
    
    latency_ms = time.time() * 1000 - start_ms
    metrics = {"latency_ms": round(latency_ms, 2)}
    
    return result, metrics


# =============================================================================
# CLI INTERFACE
# =============================================================================

def main():
    parser = argparse.ArgumentParser(
        description="Convert Microsoft Word documents to markdown",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  sft_docx2md.py read document.docx
  sft_docx2md.py read --no-tables document.docx
  sft_docx2md.py analyze document.docx
  sft_docx2md.py read document.docx > doc.md
  sft_docx2md.py read document.docx | sft_md2html.py --show
        """,
    )
    
    parser.add_argument(
        "-V", "--version",
        action="version",
        version=f"%(prog)s {CONFIG['version']}",
    )
    
    subparsers = parser.add_subparsers(dest="command", help="Commands")
    
    # read
    p_read = subparsers.add_parser("read", help="Extract text as markdown")
    p_read.add_argument("filepath", nargs="?", help="Path to .docx file (or stdin)")
    p_read.add_argument(
        "-t", "--tables",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Include/exclude tables (default: include)",
    )
    
    # analyze
    p_analyze = subparsers.add_parser("analyze", help="Analyze document structure")
    p_analyze.add_argument("filepath", help="Path to .docx file")
    
    # MCP server
    subparsers.add_parser("mcp-stdio", help="Run as MCP server")
    
    args = parser.parse_args()
    
    try:
        if args.command == "mcp-stdio":
            _run_mcp()
        elif args.command == "read":
            filepath = args.filepath
            if not filepath and not sys.stdin.isatty():
                filepath = sys.stdin.read().strip()
            assert filepath, "filepath required (positional argument or stdin)"
            result, metrics = _read_impl(filepath, args.tables)
            print(result)
            _log("INFO", "read", f"file={filepath}", metrics=json.dumps(metrics))
        elif args.command == "analyze":
            result, metrics = _analyze_impl(args.filepath)
            print(json.dumps(result, indent=2))
            _log("INFO", "analyze", f"file={args.filepath}", metrics=json.dumps(metrics))
        else:
            parser.print_help()
    except AssertionError as e:
        _log("ERROR", "contract_violation", str(e))
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)
    except Exception as e:
        _log("ERROR", "runtime_error", str(e))
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)


# =============================================================================
# FASTMCP SERVER
# =============================================================================

def _run_mcp():
    """Run as MCP server."""
    from fastmcp import FastMCP
    
    mcp = FastMCP("docx2md")
    
    @mcp.tool()
    def read(filepath: str, include_tables: bool = True) -> str:
        """Convert Word document to markdown text.
        
        Args:
            filepath: Path to .docx file
            include_tables: Whether to include tables (default: True)
        
        Returns:
            Markdown text content
        """
        try:
            result, metrics = _read_impl(filepath, include_tables)
            return json.dumps({"markdown": result, "metrics": metrics})
        except Exception as e:
            return json.dumps({"error": str(e)})
    
    @mcp.tool()
    def analyze(filepath: str) -> str:
        """Analyze Word document structure.
        
        Args:
            filepath: Path to .docx file
        
        Returns:
            JSON with document statistics and structure
        """
        try:
            result, metrics = _analyze_impl(filepath)
            return json.dumps({"analysis": result, "metrics": metrics})
        except Exception as e:
            return json.dumps({"error": str(e)})
    
    mcp.run(transport="stdio")


if __name__ == "__main__":
    main()
